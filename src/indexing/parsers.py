"""
Парсеры документов для KAG

Поддерживаемые форматы:
- PDF (с OCR для таблиц и изображений)
- TXT, MD
- DOCX
- CSV

Извлекает текст, таблицы и метаданные для последующей векторизации.
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import os
from loguru import logger


class DocumentSegment:
    """Сегмент документа (текст, таблица, изображение)"""
    
    def __init__(
        self,
        segment_type: str,
        content: str,
        page: Optional[int] = None,
        metadata: Optional[Dict[str, Any]] = None
    ):
        self.segment_type = segment_type  # text, table, image
        self.content = content
        self.page = page
        self.metadata = metadata or {}


class DocumentParser:
    """
    Парсер документов с поддержкой мультимодальных данных.
    
    Извлекает текст, метаданные и структуру из различных форматов.
    """

    def __init__(self):
        self._supported_formats = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_txt,
            ".md": self._parse_txt,
            ".docx": self._parse_docx,
            ".csv": self._parse_csv,
        }

    def parse(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Распарсить документ.

        Args:
            file_path: Путь к файлу
            file_type: MIME-тип или расширение (опционально)

        Returns:
            Словарь с содержимым и метаданными:
            {
                "segments": [...],
                "metadata": {...},
                "total_pages": N,
                "total_segments": N
            }
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        # Определяем парсер по расширению
        extension = path.suffix.lower()
        parser_func = self._supported_formats.get(extension)

        if not parser_func:
            raise ValueError(f"Неподдерживаемый формат: {extension}. Поддерживаемые: {list(self._supported_formats.keys())}")

        logger.info(f"Парсинг файла: {file_path}, формат: {extension}")

        try:
            result = parser_func(path)
            
            # Добавляем общую информацию
            file_stat = path.stat()
            result["metadata"].update({
                "filename": path.name,
                "file_path": str(path),
                "file_size": file_stat.st_size,
                "file_type": extension,
                "created_at": datetime.utcnow().isoformat(),
                "total_segments": len(result.get("segments", []))
            })

            logger.info(
                f"Документ распарсен: {path.name}, "
                f"сегментов: {result['metadata']['total_segments']}"
            )

            return result

        except Exception as e:
            logger.error(f"Ошибка парсинга {file_path}: {e}")
            raise

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """
        Распарсить PDF файл.

        Использует PyPDF2 для извлечения текста.
        Если текст на странице пустой или слишком короткий — автоматически
        применяется OCR через Tesseract (распознавание русского/английского).
        """
        try:
            import PyPDF2
        except ImportError:
            logger.warning("PyPDF2 не установлен. Установите: pip install PyPDF2")
            return self._create_fallback_result(path, "pdf", "PyPDF2 не установлен")

        segments = []
        metadata = {}
        ocr_used = False

        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                # Метаданные PDF
                if reader.metadata:
                    metadata = {
                        "author": reader.metadata.get('/Author', ''),
                        "title": reader.metadata.get('/Title', ''),
                        "subject": reader.metadata.get('/Subject', ''),
                        "creator": reader.metadata.get('/Creator', ''),
                        "producer": reader.metadata.get('/Producer', ''),
                    }

                # Извлечение текста по страницам
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()

                    # Если текст пустой или слишком короткий — используем OCR
                    if not text or len(text.strip()) < 50:
                        logger.info(
                            f"Страница {page_num + 1}: текст не найден или "
                            f"слишком короткий ({len(text) if text else 0} символов), "
                            f"применяю OCR..."
                        )
                        
                        # Сначала пробуем Tesseract
                        from src.indexing.ocr_engine import ocr_engine
                        if ocr_engine.is_available:
                            try:
                                ocr_result = ocr_engine.extract_text_from_pdf(str(path))
                                if ocr_result.get("pages") and page_num < len(ocr_result["pages"]):
                                    page_data = ocr_result["pages"][page_num]
                                    text = page_data.get("text", "")
                                    if text:
                                        ocr_used = True
                                        logger.info(
                                            f"Страница {page_num + 1}: Tesseract OCR распознал "
                                            f"{len(text)} символов"
                                        )
                            except Exception as e:
                                logger.warning(f"Tesseract OCR не работает: {e}")
                        
                        # Если Tesseract не помог - пробуем LLM OCR
                        if not text or len(text.strip()) < 50:
                            try:
                                from src.indexing.llm_ocr import get_llm_ocr_engine
                                llm_ocr = get_llm_ocr_engine()
                                if llm_ocr.is_available:
                                    # Рендерим страницу в изображение
                                    from pdf2image import convert_from_path
                                    images = convert_from_path(str(path), dpi=150, first_page=page_num+1, last_page=page_num+1)
                                    if images:
                                        import tempfile
                                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                                            images[0].save(tmp.name, "PNG")
                                            text = llm_ocr.extract_from_image(tmp.name)
                                            if text:
                                                ocr_used = True
                                                logger.info(
                                                    f"Страница {page_num + 1}: LLM OCR распознал "
                                                    f"{len(text)} символов"
                                                )
                            except Exception as e:
                                logger.debug(f"LLM OCR не работает: {e}")

                    if text and text.strip():
                        segments.append(DocumentSegment(
                            segment_type="text",
                            content=text.strip(),
                            page=page_num + 1,
                            metadata={
                                "page_number": page_num + 1,
                                "char_count": len(text),
                                "ocr_used": ocr_used
                            }
                        ).__dict__)

                metadata["total_pages"] = len(reader.pages)
                metadata["ocr_used"] = ocr_used

        except Exception as e:
            logger.error(f"Ошибка чтения PDF {path}: {e}")
            raise

        return {
            "segments": segments,
            "metadata": metadata
        }

    def _parse_txt(self, path: Path) -> Dict[str, Any]:
        """Распарсить текстовый файл"""
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # Пробуем другие кодировки
            content = path.read_text(encoding="latin-1")

        # Разбиваем на абзацы
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        segments = []
        for i, para in enumerate(paragraphs):
            segments.append({
                "segment_type": "text",
                "content": para,
                "page": 1,
                "metadata": {
                    "paragraph_index": i,
                    "char_count": len(para)
                }
            })

        return {
            "segments": segments,
            "metadata": {
                "total_pages": 1,
                "total_paragraphs": len(paragraphs)
            }
        }

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """
        Распарсить DOCX файл.

        Использует python-docx для извлечения текста и таблиц.
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx не установлен. Установите: pip install python-docx")
            return self._create_fallback_result(path, "docx", "python-docx не установлен")

        doc = Document(path)
        segments = []
        
        # Метаданные
        metadata = {}
        if doc.core_properties:
            metadata = {
                "author": doc.core_properties.author or '',
                "title": doc.core_properties.title or '',
                "subject": doc.core_properties.subject or '',
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else '',
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else ''
            }

        # Извлечение текста по параграфам
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                segments.append({
                    "segment_type": "text",
                    "content": para.text.strip(),
                    "page": 1,
                    "metadata": {
                        "paragraph_index": i,
                        "style": para.style.name if para.style else '',
                        "char_count": len(para.text)
                    }
                })

        # Извлечение таблиц
        table_count = 0
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))
            
            if table_rows:
                table_text = "\n".join(table_rows)
                segments.append({
                    "segment_type": "table",
                    "content": table_text,
                    "page": 1,
                    "metadata": {
                        "table_index": table_idx,
                        "rows": len(table.rows),
                        "columns": max(len(row.cells) for row in table.rows) if table.rows else 0,
                        "char_count": len(table_text)
                    }
                })
                table_count += 1

        metadata["total_paragraphs"] = len([s for s in segments if s.get("segment_type") == "text"])
        metadata["total_tables"] = table_count

        return {
            "segments": segments,
            "metadata": metadata
        }

    def _parse_csv(self, path: Path) -> Dict[str, Any]:
        """Распарсить CSV файл"""
        import csv
        
        rows = []
        headers = []
        
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            rows = [row for row in reader]

        # Преобразуем в текстовые сегменты
        segments = []
        for i, row in enumerate(rows):
            row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row))
            segments.append({
                "segment_type": "text",
                "content": row_text,
                "page": 1,
                "metadata": {
                    "row_index": i,
                    "columns": dict(zip(headers, row))
                }
            })

        return {
            "segments": segments,
            "metadata": {
                "total_rows": len(rows),
                "columns": headers
            }
        }

    def _create_fallback_result(self, path: Path, file_type: str, error: str) -> Dict[str, Any]:
        """Создать результат-заглушку при ошибке парсинга"""
        return {
            "segments": [{
                "segment_type": "text",
                "content": f"Не удалось распарсить {file_type} файл. Требуется установка зависимостей: {error}",
                "page": 1,
                "metadata": {"error": error}
            }],
            "metadata": {"error": error, "total_pages": 0}
        }


class TextChunker:
    """
    Разбиение текста на чанки для векторизации.
    
    Стратегии:
    - По размеру (фиксированное количество символов)
    - С перекрытием для сохранения контекста
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_document(self, segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Разбить сегменты документа на чанки.

        Args:
            segments: Список сегментов из парсера

        Returns:
            Список чанков для векторизации
        """
        chunks = []
        
        for segment in segments:
            content = segment.get("content", "")
            metadata = segment.get("metadata", {})
            
            if len(content) <= self.chunk_size:
                # Сегмент помещается в один чанк
                chunks.append({
                    "chunk_id": f"chunk_{len(chunks)}",
                    "content": content,
                    "metadata": {
                        **metadata,
                        "chunk_index": len(chunks),
                        "is_partial": False
                    }
                })
            else:
                # Разбиваем на несколько чанков
                segment_chunks = self._split_content(content, metadata)
                chunks.extend(segment_chunks)

        logger.info(f"Документ разбит на {len(chunks)} чанков")
        return chunks

    def _split_content(
        self,
        content: str,
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Разбить длинный текст на чанки с перекрытием"""
        
        chunks = []
        start = 0
        chunk_index = 0

        while start < len(content):
            end = start + self.chunk_size

            # Пытаемся разбить по границе слова
            if end < len(content):
                space_pos = content.rfind(" ", start + self.chunk_size // 2, end)
                if space_pos > start:
                    end = space_pos + 1

            chunk_text = content[start:end].strip()

            if chunk_text:
                chunks.append({
                    "chunk_id": f"chunk_{len(chunks)}",
                    "content": chunk_text,
                    "metadata": {
                        **metadata,
                        "chunk_index": chunk_index,
                        "start_pos": start,
                        "end_pos": end,
                        "char_count": len(chunk_text),
                        "is_partial": True
                    }
                })
                chunk_index += 1

            start = end - self.chunk_overlap

        return chunks


# Глобальные экземпляры
document_parser = DocumentParser()
text_chunker = TextChunker(chunk_size=1000, chunk_overlap=200)
