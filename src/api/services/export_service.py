"""
Document Export Service для KAG

Экспорт диалогов чата в форматы:
- DOCX (Word документ)
- PDF (PDF документ)

Использует python-docx и reportlab для прямой генерации без внешних зависимостей.
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import io
from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx не установлен")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("reportlab не установлен")


class DocumentExportService:
    """
    Сервис экспорта диалогов чата в документы.
    
    Поддерживаемые форматы:
    - DOCX (Word)
    - PDF
    """

    def __init__(self):
        """Инициализация сервиса"""
        logger.info("DocumentExportService инициализирован")

    def export_to_docx(
        self,
        messages: List[Dict[str, Any]],
        title: str = "Диалог с KAG",
        author: str = "KAG System",
        include_metadata: bool = True
    ) -> bytes:
        """
        Экспортировать диалог в DOCX.

        Args:
            messages: Список сообщений [{"role": "...", "content": "...", ...}]
            title: Заголовок документа
            author: Автор документа
            include_metadata: Включать ли метаданные (модель, время и т.д.)

        Returns:
            Байты DOCX файла
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx не установлен. pip install python-docx")

        doc = Document()
        
        # Настройка стилей
        self._setup_docx_styles(doc)
        
        # Титульная страница
        self._add_docx_title(doc, title, author)
        doc.add_page_break()
        
        # Оглавление
        doc.add_heading('Содержание', level=1)
        doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Количество сообщений: {len(messages)}')
        doc.add_paragraph(f'Автор: {author}')
        doc.add_page_break()
        
        # Диалог
        doc.add_heading('Диалог', level=1)
        
        for i, msg in enumerate(messages, 1):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            metadata = msg.get("metadata", {})
            
            if role == "user":
                self._add_docx_user_message(doc, content, i, metadata if include_metadata else None)
            elif role == "assistant":
                self._add_docx_assistant_message(doc, content, i, metadata if include_metadata else None)
        
        # Сохраняем в байты
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"DOCX документ создан: {len(messages)} сообщений")
        return buffer.getvalue()

    def export_to_pdf(
        self,
        messages: List[Dict[str, Any]],
        title: str = "Диалог с KAG",
        author: str = "KAG System",
        include_metadata: bool = True
    ) -> bytes:
        """
        Экспортировать диалог в PDF.

        Args:
            messages: Список сообщений
            title: Заголовок документа
            author: Автор
            include_metadata: Включать метаданные

        Returns:
            Байты PDF файла
        """
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab не установлен. pip install reportlab")

        buffer = io.BytesIO()
        
        # Создаём документ
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Стили
        styles = self._setup_pdf_styles()
        
        # Содержимое
        story = []
        
        # Заголовок
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}', styles['Normal']))
        story.append(Paragraph(f'Автор: {author}', styles['Normal']))
        story.append(Paragraph(f'Сообщений: {len(messages)}', styles['Normal']))
        story.append(Spacer(1, 1*cm))
        
        # Разделитель
        story.append(Paragraph('_' * 80, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
        # Диалог
        for i, msg in enumerate(messages, 1):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            metadata = msg.get("metadata", {})
            
            if role == "user":
                self._add_pdf_user_message(story, styles, content, i, metadata if include_metadata else None)
            elif role == "assistant":
                self._add_pdf_assistant_message(story, styles, content, i, metadata if include_metadata else None)
        
        # Строим PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info(f"PDF документ создан: {len(messages)} сообщений")
        return buffer.getvalue()

    # ===========================================
    # DOCX методы
    # ===========================================

    def _setup_docx_styles(self, doc):
        """Настройка стилей DOCX"""
        pass  # Используем стандартные стили

    def _add_docx_title(self, doc, title, author):
        """Добавить титульную страницу"""
        for _ in range(6):
            doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        
        doc.add_paragraph()
        
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f'Создано: {author}')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    def _add_docx_user_message(self, doc, content, num, metadata=None):
        """Добавить сообщение пользователя"""
        p = doc.add_paragraph()
        run = p.add_run(f'Пользователь (#{num})')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        
        doc.add_paragraph(content)
        
        if metadata:
            self._add_docx_metadata(doc, metadata)

    def _add_docx_assistant_message(self, doc, content, num, metadata=None):
        """Добавить сообщение ассистента"""
        p = doc.add_paragraph()
        run = p.add_run(f'Ассистент (#{num})')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
        
        doc.add_paragraph(content)
        
        if metadata:
            self._add_docx_metadata(doc, metadata)

    def _add_docx_metadata(self, doc, metadata):
        """Добавить метаданные"""
        meta_parts = []
        if 'model' in metadata:
            meta_parts.append(f"Модель: {metadata['model']}")
        if 'backend' in metadata:
            meta_parts.append(f"Бэкенд: {metadata['backend']}")
        
        if meta_parts:
            p = doc.add_paragraph()
            run = p.add_run(' | '.join(meta_parts))
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ===========================================
    # PDF методы
    # ===========================================

    def _setup_pdf_styles(self):
        """Создание стилей для PDF"""
        styles = getSampleStyleSheet()
        
        # Заголовок
        styles.add(ParagraphStyle(
            name='KAGTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=HexColor('#3B82F6'),
            spaceAfter=20,
            alignment=TA_CENTER
        ))
        
        # Имя пользователя
        styles.add(ParagraphStyle(
            name='UserName',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=HexColor('#1E40AF'),
            spaceBefore=15,
            spaceAfter=5
        ))
        
        # Имя ассистента
        styles.add(ParagraphStyle(
            name='AssistantName',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=HexColor('#8B5CF6'),
            spaceBefore=15,
            spaceAfter=5
        ))
        
        # Текст сообщения
        styles.add(ParagraphStyle(
            name='MessageText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=10
        ))
        
        # Метаданные
        styles.add(ParagraphStyle(
            name='MetaText',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('#64748B'),
            fontName='Helvetica-Oblique',
            spaceAfter=15
        ))
        
        return styles

    def _add_pdf_user_message(self, story, styles, content, num, metadata=None):
        """Добавить сообщение пользователя в PDF"""
        story.append(Paragraph(f'👤 Пользователь (#{num})', styles['UserName']))
        
        # Экранируем HTML
        safe_content = self._escape_html(content)
        story.append(Paragraph(safe_content, styles['MessageText']))
        
        if metadata:
            self._add_pdf_metadata(story, styles, metadata)

    def _add_pdf_assistant_message(self, story, styles, content, num, metadata=None):
        """Добавить сообщение ассистента в PDF"""
        story.append(Paragraph(f'🤖 Ассистент (#{num})', styles['AssistantName']))
        
        safe_content = self._escape_html(content)
        story.append(Paragraph(safe_content, styles['MessageText']))
        
        if metadata:
            self._add_pdf_metadata(story, styles, metadata)

    def _add_pdf_metadata(self, story, styles, metadata):
        """Добавить метаданные в PDF"""
        meta_parts = []
        if 'model' in metadata:
            meta_parts.append(f"Модель: {metadata['model']}")
        if 'backend' in metadata:
            meta_parts.append(f"Бэкенд: {metadata['backend']}")
        
        if meta_parts:
            story.append(Paragraph(' | '.join(meta_parts), styles['MetaText']))

    def _escape_html(self, text: str) -> str:
        """Экранировать текст для PDF"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('\n', '<br/>'))


# Глобальный экземпляр
export_service = DocumentExportService()
